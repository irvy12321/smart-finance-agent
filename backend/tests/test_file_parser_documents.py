"""
PDF / DOCX 正向解析测试（真实字节 → 真实解析器）

之前只有"依赖未安装时报错"的负向分支（且在依赖已装的环境永远 skip），
PDF/DOCX 的真实提取路径没有任何覆盖。本文件在内存中构造真实的
PDF/DOCX 字节流，验证 parse_file 能提取出正确文本。
"""

import io

import pytest

from app.rag.file_parser import (
    FileParserError,
    parse_docx_file,
    parse_file,
    parse_pdf_file,
)


def _build_minimal_pdf(text: str) -> bytes:
    """构造一个带正确 xref 表的单页 PDF，页面内容为给定文本"""
    stream = f"BT /F1 24 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length "
        + str(len(stream)).encode()
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f"{i} 0 obj\n".encode())
        out.write(obj)
        out.write(b"\nendobj\n")

    xref_pos = out.tell()
    out.write(f"xref\n0 {len(objects) + 1}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for off in offsets:
        out.write(f"{off:010d} 00000 n \n".encode())
    out.write(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n".encode()
    )
    return out.getvalue()


def _build_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table:
        t = doc.add_table(rows=len(table), cols=len(table[0]))
        for r, row in enumerate(table):
            for c, cell in enumerate(row):
                t.cell(r, c).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestPdfParsing:
    def test_parse_real_pdf_extracts_text(self):
        content = _build_minimal_pdf("Quarterly revenue grew 12 percent")
        text = parse_pdf_file(content, "report.pdf")
        assert "Quarterly revenue grew 12 percent" in text

    def test_parse_file_dispatches_pdf(self):
        content = _build_minimal_pdf("Hello PDF")
        text = parse_file(content, "doc.PDF")
        assert "Hello PDF" in text

    def test_corrupted_pdf_raises(self):
        with pytest.raises(FileParserError):
            parse_pdf_file(b"not a pdf at all", "bad.pdf")


class TestDocxParsing:
    def test_parse_real_docx_extracts_paragraphs(self):
        content = _build_docx(["First paragraph.", "Second paragraph."])
        text = parse_docx_file(content, "memo.docx")
        assert "First paragraph." in text
        assert "Second paragraph." in text

    def test_parse_docx_extracts_table_cells(self):
        content = _build_docx(
            ["Financials"], table=[["Metric", "Value"], ["Revenue", "100M"]]
        )
        text = parse_docx_file(content, "table.docx")
        assert "Metric | Value" in text
        assert "Revenue | 100M" in text

    def test_parse_file_dispatches_docx(self):
        content = _build_docx(["Dispatch check"])
        text = parse_file(content, "memo.DOCX")
        assert "Dispatch check" in text

    def test_empty_docx_raises(self):
        content = _build_docx([])
        with pytest.raises(FileParserError):
            parse_docx_file(content, "empty.docx")

    def test_corrupted_docx_raises(self):
        with pytest.raises(FileParserError):
            parse_docx_file(b"not a docx", "bad.docx")
