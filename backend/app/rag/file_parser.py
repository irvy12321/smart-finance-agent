"""
文件解析器 - 支持多种文件格式的文本提取

支持格式:
- .txt, .md, .csv, .json: UTF-8 文本
- .pdf: PyPDF2 提取
- .docx: python-docx 提取
"""
import json
from pathlib import Path

from app.utils.logger import get_logger

logger = get_logger("file_parser")


class FileParserError(Exception):
    """文件解析错误"""
    pass


def parse_text_file(content: bytes, filename: str) -> str:
    """解析文本文件 (txt, md, csv, json)"""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        # 尝试其他编码
        try:
            return content.decode("gbk")
        except UnicodeDecodeError:
            try:
                return content.decode("latin-1")
            except Exception as e:
                raise FileParserError(f"无法解码文件 {filename}: {e}") from e


def parse_pdf_file(content: bytes, filename: str) -> str:
    """解析 PDF 文件"""
    try:
        from PyPDF2 import PdfReader
        import io
        
        pdf_reader = PdfReader(io.BytesIO(content))
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"PDF 第 {page_num + 1} 页提取失败: {e}")
                continue
        
        if not text_parts:
            raise FileParserError(f"PDF 文件 {filename} 无法提取任何文本")
        
        return "\n\n".join(text_parts)
    except ImportError:
        raise FileParserError("PyPDF2 未安装，无法解析 PDF 文件")
    except Exception as e:
        raise FileParserError(f"PDF 解析失败: {e}") from e


def parse_docx_file(content: bytes, filename: str) -> str:
    """解析 DOCX 文件"""
    try:
        from docx import Document
        import io
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            raise FileParserError(f"DOCX 文件 {filename} 无法提取任何文本")
        
        return "\n\n".join(text_parts)
    except ImportError:
        raise FileParserError("python-docx 未安装，无法解析 DOCX 文件")
    except Exception as e:
        raise FileParserError(f"DOCX 解析失败: {e}") from e


def parse_file(content: bytes, filename: str) -> str:
    """
    根据文件类型自动选择解析器
    
    Args:
        content: 文件内容 (bytes)
        filename: 文件名
    
    Returns:
        提取的文本内容
    
    Raises:
        FileParserError: 解析失败
    """
    file_ext = Path(filename).suffix.lower()
    
    # 文本文件
    text_extensions = {".txt", ".md", ".csv", ".json"}
    if file_ext in text_extensions:
        logger.info(f"解析文本文件: {filename}")
        return parse_text_file(content, filename)
    
    # PDF 文件
    if file_ext == ".pdf":
        logger.info(f"解析 PDF 文件: {filename}")
        return parse_pdf_file(content, filename)
    
    # DOCX 文件
    if file_ext == ".docx":
        logger.info(f"解析 DOCX 文件: {filename}")
        return parse_docx_file(content, filename)
    
    raise FileParserError(f"不支持的文件类型: {file_ext}")


def get_supported_extensions() -> set[str]:
    """获取支持的文件扩展名"""
    return {".txt", ".md", ".csv", ".json", ".pdf", ".docx"}
